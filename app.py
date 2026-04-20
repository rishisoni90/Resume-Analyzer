from flask import Flask, request, jsonify, render_template_string, send_file
import os
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from nltk.stem import WordNetLemmatizer
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                 HRFlowable, Table, TableStyle, KeepTogether)
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import io

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    for pkg in ('punkt','stopwords','wordnet','averaged_perceptron_tagger'):
        nltk.download(pkg, quiet=True)

app = Flask(__name__)
app.config['OPTIMIZED_FOLDER'] = 'optimized_resumes'
os.makedirs(app.config['OPTIMIZED_FOLDER'], exist_ok=True)

# ──────────────────────────────────────────────────────────────
# SKILL SYNONYMS & WEIGHTS
# ──────────────────────────────────────────────────────────────
SKILL_SYNONYMS = {
    'javascript': ['js','nodejs','react','vue','angular'],
    'python':     ['django','flask','fastapi','pandas','numpy'],
    'java':       ['spring','springboot','hibernate'],
    'aws':        ['ec2','s3','lambda','cloudformation','dynamodb'],
    'docker':     ['containers','dockerfile','docker-compose'],
    'kubernetes': ['k8s','helm','openshift'],
    'sql':        ['mysql','postgresql','postgres','mongodb','nosql'],
    'git':        ['github','gitlab','bitbucket'],
    'ci/cd':      ['jenkins','github actions','circleci'],
    'agile':      ['scrum','kanban','jira'],
    'cloud':      ['azure','gcp','google cloud','serverless'],
    'devops':     ['sre','terraform','ansible','infrastructure'],
    'testing':    ['qa','selenium','unit testing'],
    'machine learning': ['ml','ai','tensorflow','pytorch'],
}

HIGH_WEIGHT   = {'architecture','lead','senior','scalability','performance','security',
                 'distributed','microservices','api','cloud','aws','azure','gcp'}
MEDIUM_WEIGHT = {'agile','scrum','ci/cd','docker','kubernetes','deployment','testing'}

# ──────────────────────────────────────────────────────────────
# MATCHER
# ──────────────────────────────────────────────────────────────
class ResumeMatcher:
    def __init__(self):
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        self.impact_verbs = [
            'increased','decreased','reduced','improved','boosted','saved','generated',
            'delivered','achieved','accelerated','optimized','enhanced','streamlined',
            'led','built','launched','deployed','automated','engineered','designed',
            'cut','grew','expanded','maximized','minimized',
        ]

    def preprocess(self, text):
        text = re.sub(r'[^a-zA-Z\s]', '', text.lower())
        return [self.lemmatizer.lemmatize(w)
                for w in word_tokenize(text)
                if w not in self.stop_words and len(w) > 2]

    def extract_skills(self, text):
        tl = text.lower()
        skills = set(self.preprocess(text))
        for skill, syns in SKILL_SYNONYMS.items():
            if skill in tl or any(s in tl for s in syns):
                skills.add(skill); skills.update(syns)
        return skills

    def impact_analysis(self, text):
        found = []
        for sent in re.split(r'[.\n!?]+', text):
            if any(v in sent.lower() for v in self.impact_verbs) and re.search(r'\d+', sent):
                found.append(sent.strip())
        pct  = len(re.findall(r'\d+%', text))
        curr = len(re.findall(r'\$\d+', text))
        verbs= sum(1 for v in self.impact_verbs if v in text.lower())
        score = min(100, len(set(found))*8 + pct*6 + curr*4 + verbs*2)
        suggestions = []
        if pct < 2:
            suggestions.append("Add percentage improvements e.g. 'Reduced errors by 40%'")
        if curr == 0:
            suggestions.append("Include dollar amounts e.g. 'Saved $50,000 annually'")
        if verbs < 5:
            suggestions.append("Use strong action verbs: Led, Architected, Optimized, Delivered")
        if len(set(found)) < 3:
            suggestions.append("Add time-based wins e.g. 'Cut processing time from 4h to 30 min'")
        if not suggestions:
            suggestions.append("Strong quantified impacts — great job!")
        return {'score': score, 'total': len(set(found)),
                'needs': score < 60, 'suggestions': suggestions}

    def match_score(self, resume, job):
        rs, js = self.extract_skills(resume), self.extract_skills(job)
        if not js: return 0, set(), set(), {}
        hits = rs & js
        wm = wt = 0
        for k in js:
            w = 1.5 if k in HIGH_WEIGHT else (1.2 if k in MEDIUM_WEIGHT else 1.0)
            wt += w
            if k in hits: wm += w
        pct = (wm/wt*100) if wt else 0
        ia  = self.impact_analysis(resume)
        exp = 5 if (re.findall(r'(\d+)\s*years?', resume, re.I) and
                    re.findall(r'(\d+)\s*years?', job, re.I) and
                    int(max(re.findall(r'(\d+)\s*years?', resume, re.I))) >=
                    int(min(re.findall(r'(\d+)\s*years?', job, re.I)))) else 0
        return min(100, pct + ia['score']*0.1 + exp), hits, js-rs, ia

    def extract_name(self, text, file_fmt=False):
        for line in text.strip().split('\n')[:5]:
            line = line.strip()
            if not line: continue
            if any(x in line.lower() for x in ['@','linkedin','github','+','|','http']): continue
            if 2 < len(line) < 55 and re.match(r'^[A-Za-z\s\.]+$', line):
                return re.sub(r'\s+','_',line) if file_fmt else line
        return 'Candidate'

    def extract_job_role(self, job_text):
        for line in job_text.split('\n')[:10]:
            for role in ['Software Engineer','DevOps Engineer','Site Reliability Engineer',
                         'Data Scientist','Full Stack Developer','Machine Learning Engineer',
                         'Cloud Engineer','Backend Developer','Frontend Developer',
                         'Engineering Manager','Technical Lead','Architect']:
                if role.lower() in line.lower():
                    return role.replace(' ','_')
        return 'Position'

    def add_keywords(self, text, missing):
        if not missing: return text
        kws = ', '.join(k.title() for k in list(missing)[:10])
        lines, added = text.split('\n'), False
        out = []
        for l in lines:
            out.append(l)
            if not added and re.search(r'SKILLS|TECHNICAL', l.upper()):
                out.append(f'  {kws}'); added = True
        if not added: out.insert(3, f'\nSKILLS\n{kws}\n')
        return '\n'.join(out)

    # ──────────────────────────────────────────────────────────
    # PDF  — fully redesigned, clean, professional
    # ──────────────────────────────────────────────────────────
    def build_pdf(self, resume_text, output_path):
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=letter,
            leftMargin=0.65*inch, rightMargin=0.65*inch,
            topMargin=0.6*inch,  bottomMargin=0.6*inch)

        # colour palette
        C_DARK  = colors.HexColor('#111827')
        C_BLUE  = colors.HexColor('#1D4ED8')
        C_GRAY  = colors.HexColor('#374151')
        C_LGRAY = colors.HexColor('#6B7280')
        C_GREEN = colors.HexColor('#14532D')
        C_WHITE = colors.white

        W = 7.25 * inch   # usable width

        # ── styles ──────────────────────────────────────────
        name_sty = ParagraphStyle('N', fontSize=26, fontName='Helvetica-Bold',
            textColor=C_DARK, alignment=TA_CENTER, leading=30, spaceAfter=2)

        contact_sty = ParagraphStyle('C', fontSize=9, fontName='Helvetica',
            textColor=C_LGRAY, alignment=TA_CENTER, leading=13, spaceAfter=8)

        sec_sty = ParagraphStyle('S', fontSize=9, fontName='Helvetica-Bold',
            textColor=C_BLUE, leading=11, spaceBefore=0, spaceAfter=0)

        job_title_sty = ParagraphStyle('JT', fontSize=10.5, fontName='Helvetica-Bold',
            textColor=C_DARK, leading=14, spaceBefore=10, spaceAfter=1)

        job_meta_sty = ParagraphStyle('JM', fontSize=9, fontName='Helvetica',
            textColor=C_LGRAY, leading=12, spaceAfter=4)

        bullet_sty = ParagraphStyle('B', fontSize=9.5, fontName='Helvetica',
            textColor=C_GRAY, leading=13, leftIndent=14,
            firstLineIndent=-10, spaceBefore=2, spaceAfter=2)

        bullet_hit_sty = ParagraphStyle('BH', fontSize=9.5, fontName='Helvetica-Bold',
            textColor=C_GREEN, leading=13, leftIndent=14,
            firstLineIndent=-10, spaceBefore=2, spaceAfter=2)

        skill_sty = ParagraphStyle('SK', fontSize=9.5, fontName='Helvetica',
            textColor=C_GRAY, leading=13, spaceBefore=2, spaceAfter=2)

        body_sty = ParagraphStyle('BD', fontSize=9.5, fontName='Helvetica',
            textColor=C_GRAY, leading=14, spaceBefore=2, spaceAfter=2)

        proj_url_sty = ParagraphStyle('PU', fontSize=8.5, fontName='Helvetica',
            textColor=C_BLUE, leading=12, spaceAfter=3)

        # ── section header helper ───────────────────────────
        def section_header(title):
            tbl = Table([[Paragraph(title.upper(), sec_sty)]], colWidths=[W])
            tbl.setStyle(TableStyle([
                ('TOPPADDING',    (0,0),(-1,-1), 4),
                ('BOTTOMPADDING', (0,0),(-1,-1), 4),
                ('LEFTPADDING',   (0,0),(-1,-1), 0),
                ('LINEBELOW',     (0,0),(-1,-1), 1.2, C_BLUE),
            ]))
            return [Spacer(1, 0.1*inch), tbl, Spacer(1, 0.06*inch)]

        # ── job block helper ────────────────────────────────
        def job_block(line):
            parts = [p.strip() for p in line.split('|')]
            items = [Paragraph(parts[0], job_title_sty)]
            if len(parts) > 1:
                items.append(Paragraph('  ·  '.join(parts[1:]), job_meta_sty))
            return items

        # ── parse & build story ─────────────────────────────
        story = []
        name_done = contact_done = False
        lines = resume_text.strip().split('\n')

        for line in lines:
            ls = line.strip()

            # blank
            if not ls:
                story.append(Spacer(1, 0.04*inch))
                continue

            # NAME
            if not name_done and len(ls) < 65 and \
               not any(x in ls for x in ['@','|','http','+']):
                story.append(Paragraph(ls, name_sty))
                story.append(HRFlowable(width='100%', thickness=2,
                    color=C_BLUE, spaceBefore=4, spaceAfter=6))
                name_done = True
                continue

            # CONTACT
            if not contact_done and any(x in ls.lower() for x in ['@','linkedin','github','+']):
                story.append(Paragraph(ls, contact_sty))
                contact_done = True
                continue

            # SECTION HEADER
            if ls.isupper() and 2 < len(ls) < 55:
                story += section_header(ls)
                continue

            # JOB / ROLE LINE  (contains | but isn't a bullet)
            if '|' in ls and len(ls) < 140 and not ls.startswith(('•','-','*')):
                story.append(KeepTogether(job_block(ls)))
                continue

            # PROJECT GITHUB URL
            if ls.lower().startswith(('github.com','http','www.')):
                story.append(Paragraph(ls, proj_url_sty))
                continue

            # SKILL LINE  "Label: values"
            if ':' in ls and not ls.startswith(('•','-','*')) and len(ls) < 140:
                idx = ls.index(':')
                key, val = ls[:idx].strip(), ls[idx+1:].strip()
                story.append(Paragraph(f'<b>{key}:</b>  {val}', skill_sty))
                continue

            # BULLET
            if ls.startswith(('•','-','*','–','▸')):
                txt = ls[1:].strip()
                has_metric = bool(re.search(
                    r'\d+%|\$[\d,]+|\d+\s*(years?|months?|hours?|x\b|times?)',
                    txt, re.I))
                sty = bullet_hit_sty if has_metric else bullet_sty
                story.append(Paragraph(f'▸  {txt}', sty))
                continue

            # DEFAULT
            story.append(Paragraph(ls, body_sty))

        doc.build(story)
        with open(output_path, 'wb') as f:
            f.write(buf.getvalue())

    # ──────────────────────────────────────────────────────────
    # WORD DOC — clean, well aligned
    # ──────────────────────────────────────────────────────────
    def build_word(self, resume_text, output_path):
        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Inches(0.70)
            sec.bottom_margin = Inches(0.70)
            sec.left_margin   = Inches(0.75)
            sec.right_margin  = Inches(0.75)

        C_DARK  = RGBColor(17,  24,  39)
        C_BLUE  = RGBColor(29,  78, 216)
        C_GRAY  = RGBColor(55,  65,  81)
        C_LGRAY = RGBColor(107,114,128)
        C_GREEN = RGBColor(20,  83,  45)
        C_WHITE = RGBColor(255,255,255)

        def set_shading(paragraph, fill_hex):
            pPr = paragraph._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'),  fill_hex)
            pPr.append(shd)

        name_done = contact_done = False
        lines = resume_text.strip().split('\n')

        for line in lines:
            ls = line.strip()

            if not ls:
                doc.add_paragraph()
                continue

            # NAME
            if not name_done and len(ls) < 65 and \
               not any(x in ls for x in ['@','|','http','+']):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(ls)
                r.font.name = 'Calibri Light'
                r.font.size = Pt(26)
                r.font.bold = True
                r.font.color.rgb = C_DARK
                p.paragraph_format.space_after = Pt(2)

                # blue rule
                rule = doc.add_paragraph()
                rule_r = rule.add_run()
                rule_r.font.size = Pt(1)
                from docx.oxml import OxmlElement as OE
                pPr2 = rule._p.get_or_add_pPr()
                pBdr = OE('w:pBdr')
                bottom = OE('w:bottom')
                bottom.set(qn('w:val'),   'single')
                bottom.set(qn('w:sz'),    '12')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), '1D4ED8')
                pBdr.append(bottom)
                pPr2.append(pBdr)
                rule.paragraph_format.space_after = Pt(6)
                name_done = True
                continue

            # CONTACT
            if not contact_done and any(x in ls.lower() for x in ['@','linkedin','github','+']):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(ls)
                r.font.name = 'Calibri'
                r.font.size = Pt(9)
                r.font.color.rgb = C_LGRAY
                p.paragraph_format.space_after = Pt(8)
                contact_done = True
                continue

            # SECTION HEADER — blue background
            if ls.isupper() and 2 < len(ls) < 55:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after  = Pt(6)
                r = p.add_run(f'  {ls}  ')
                r.font.name  = 'Calibri'
                r.font.size  = Pt(9.5)
                r.font.bold  = True
                r.font.color.rgb = C_WHITE
                set_shading(p, '1D4ED8')
                continue

            # JOB / ROLE LINE
            if '|' in ls and len(ls) < 140 and not ls.startswith(('•','-','*')):
                parts = [p2.strip() for p2 in ls.split('|')]
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after  = Pt(1)
                r = p.add_run(parts[0])
                r.font.name  = 'Calibri'
                r.font.size  = Pt(10.5)
                r.font.bold  = True
                r.font.color.rgb = C_DARK
                if len(parts) > 1:
                    mp = doc.add_paragraph()
                    mp.paragraph_format.space_before = Pt(0)
                    mp.paragraph_format.space_after  = Pt(4)
                    mr = mp.add_run('  ·  '.join(parts[1:]))
                    mr.font.name  = 'Calibri'
                    mr.font.size  = Pt(9)
                    mr.font.color.rgb = C_LGRAY
                continue

            # PROJECT URL
            if ls.lower().startswith(('github.com','http','www.')):
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                r = p.add_run(ls)
                r.font.name  = 'Calibri'
                r.font.size  = Pt(8.5)
                r.font.color.rgb = C_BLUE
                continue

            # SKILL LINE
            if ':' in ls and not ls.startswith(('•','-','*')) and len(ls) < 140:
                idx = ls.index(':')
                key, val = ls[:idx].strip(), ls[idx+1:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                r1 = p.add_run(key + ':  ')
                r1.font.name  = 'Calibri'
                r1.font.size  = Pt(9.5)
                r1.font.bold  = True
                r1.font.color.rgb = C_DARK
                r2 = p.add_run(val)
                r2.font.name  = 'Calibri'
                r2.font.size  = Pt(9.5)
                r2.font.color.rgb = C_GRAY
                continue

            # BULLET
            if ls.startswith(('•','-','*','–','▸')):
                txt = ls[1:].strip()
                has_metric = bool(re.search(
                    r'\d+%|\$[\d,]+|\d+\s*(years?|months?|hours?|x\b)',
                    txt, re.I))
                p = doc.add_paragraph()
                p.paragraph_format.left_indent   = Inches(0.22)
                p.paragraph_format.space_before  = Pt(2)
                p.paragraph_format.space_after   = Pt(2)
                p.paragraph_format.line_spacing  = 1.2
                r = p.add_run(f'▸  {txt}')
                r.font.name = 'Calibri'
                r.font.size = Pt(10)
                if has_metric:
                    r.font.color.rgb = C_GREEN
                    r.font.bold = True
                else:
                    r.font.color.rgb = C_GRAY
                continue

            # DEFAULT
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(ls)
            r.font.name  = 'Calibri'
            r.font.size  = Pt(10)
            r.font.color.rgb = C_GRAY

        doc.save(output_path)

    # ──────────────────────────────────────────────────────────
    # MAIN ANALYZE
    # ──────────────────────────────────────────────────────────
    def analyze(self, resume_text, job_text):
        score, hits, missing, ia = self.match_score(resume_text, job_text)
        optimized = self.add_keywords(resume_text, missing)
        base  = (self.extract_name(resume_text, file_fmt=True) + '_' +
                 self.extract_job_role(job_text) + '_' +
                 datetime.now().strftime('%Y%m%d'))
        base  = re.sub(r'[<>:"/\\|?*]','', re.sub(r'_+','_', base))

        docx_path = os.path.join(app.config['OPTIMIZED_FOLDER'], base+'.docx')
        pdf_path  = os.path.join(app.config['OPTIMIZED_FOLDER'], base+'.pdf')
        self.build_word(optimized, docx_path)
        self.build_pdf(optimized,  pdf_path)

        sc = round(score, 1)
        if   sc >= 85: verdict, msg = 'excellent', 'Excellent match! Your resume is strongly aligned.'
        elif sc >= 70: verdict, msg = 'good',      'Good match! A few improvements could make you stronger.'
        elif sc >= 55: verdict, msg = 'moderate',  'Moderate match — add suggested skills and metrics.'
        else:          verdict, msg = 'poor',      'Needs improvement — focus on missing skills and quantified wins.'

        kw_hints = []
        tech = [k for k in list(missing)[:15] if k not in {'communication','leadership','teamwork'}]
        soft = [k for k in list(missing)[:15] if k in {'communication','leadership','teamwork','collaboration'}]
        if tech: kw_hints.append(f"Technical skills to add: {', '.join(tech[:7])}")
        if soft: kw_hints.append(f"Soft skills to highlight: {', '.join(soft[:4])}")

        return {
            'match_score':        sc,
            'impact_score':       ia['score'],
            'matching_keywords':  list(hits)[:25],
            'missing_keywords':   list(missing)[:25],
            'total_matching':     len(hits),
            'total_missing':      len(missing),
            'total_impacts':      ia['total'],
            'impact_suggestions': ia['suggestions'][:4],
            'keyword_suggestions':kw_hints[:3],
            'verdict':            verdict,
            'verdict_message':    msg,
            'filename':           base+'.docx',
            'pdf_filename':       base+'.pdf',
            'name':               self.extract_name(resume_text),
        }


matcher = ResumeMatcher()

# ──────────────────────────────────────────────────────────────
# HTML UI
# ──────────────────────────────────────────────────────────────
HTML = '''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Resume Optimizer</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --bg:#0d1117;--surface:#161b22;--border:#30363d;
  --blue:#2f81f7;--green:#3fb950;--yellow:#d29922;--red:#f85149;
  --text:#e6edf3;--muted:#8b949e;--r:12px;
}
body{font-family:'Inter',sans-serif;background:var(--bg);color:var(--text);
     min-height:100vh;padding:36px 20px 72px}

/* header */
.hdr{text-align:center;margin-bottom:40px}
.badge{display:inline-block;background:linear-gradient(135deg,var(--blue),#8957e5);
  color:#fff;font-size:11px;font-weight:600;letter-spacing:1px;text-transform:uppercase;
  padding:4px 14px;border-radius:20px;margin-bottom:14px}
.hdr h1{font-size:clamp(1.9rem,4vw,2.7rem);font-weight:700;
  background:linear-gradient(135deg,#e6edf3,var(--blue));
  -webkit-background-clip:text;-webkit-text-fill-color:transparent;margin-bottom:8px}
.hdr p{color:var(--muted);font-size:14px}

/* grid */
.wrap{max-width:1280px;margin:0 auto}
.grid{display:grid;grid-template-columns:1fr 1fr;gap:18px;margin-bottom:18px}
.card{background:var(--surface);border:1px solid var(--border);border-radius:var(--r);padding:18px;display:flex;flex-direction:column}
.card-lbl{font-size:11px;font-weight:600;color:var(--muted);text-transform:uppercase;letter-spacing:.8px;margin-bottom:10px}
textarea{flex:1;min-height:440px;background:var(--bg);border:1px solid var(--border);
  border-radius:8px;color:var(--text);font-family:'Inter',monospace;font-size:12.5px;
  line-height:1.6;padding:14px;resize:vertical;transition:border-color .2s}
textarea:focus{outline:none;border-color:var(--blue);box-shadow:0 0 0 3px rgba(47,129,247,.15)}
textarea::placeholder{color:#484f58}

/* button */
.btn{width:100%;padding:15px;background:linear-gradient(135deg,var(--blue),#8957e5);
  color:#fff;font-family:'Inter',sans-serif;font-size:15px;font-weight:600;
  border:none;border-radius:var(--r);cursor:pointer;transition:opacity .2s,transform .1s;letter-spacing:.2px}
.btn:hover{opacity:.9;transform:translateY(-1px)}
.btn:active{transform:translateY(0)}
.btn:disabled{opacity:.45;cursor:not-allowed;transform:none}

/* spinner */
.spin-wrap{display:none;text-align:center;color:var(--muted);font-size:13px;margin:10px 0}
.spin-wrap.on{display:block}
@keyframes spin{to{transform:rotate(360deg)}}
.spin-icon{display:inline-block;width:14px;height:14px;border:2px solid var(--border);
  border-top-color:var(--blue);border-radius:50%;animation:spin .7s linear infinite;
  vertical-align:middle;margin-right:6px}

/* results */
#results{display:none;animation:up .4s ease}
@keyframes up{from{opacity:0;transform:translateY(14px)}to{opacity:1;transform:none}}

/* score hero */
.hero{background:var(--surface);border:1px solid var(--border);border-radius:var(--r);
  padding:36px 20px;text-align:center;margin-bottom:16px;position:relative;overflow:hidden}
.hero::before{content:'';position:absolute;inset:0;
  background:radial-gradient(ellipse at 50% 0%,rgba(47,129,247,.10) 0%,transparent 70%);
  pointer-events:none}
.score-num{font-size:68px;font-weight:700;line-height:1;margin-bottom:6px}
.score-lbl{font-size:18px;font-weight:600;margin-bottom:10px}
.pill{display:inline-block;padding:6px 18px;border-radius:30px;font-size:13px;
  background:rgba(255,255,255,.07);color:var(--muted)}

.c-excellent{color:var(--green)}
.c-good{color:#56d364}
.c-moderate{color:var(--yellow)}
.c-poor{color:var(--red)}

/* stat row */
.stats{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:16px}
.stat{background:var(--surface);border:1px solid var(--border);border-radius:var(--r);
  padding:16px;text-align:center}
.stat-v{font-size:28px;font-weight:700;color:var(--blue);margin-bottom:4px}
.stat-l{font-size:11px;color:var(--muted)}

/* panels */
.panel{background:var(--surface);border:1px solid var(--border);border-radius:var(--r);
  padding:16px 18px;margin-bottom:12px}
.panel-t{font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:.7px;
  color:var(--muted);margin-bottom:10px}
.tag{display:inline-block;padding:4px 10px;border-radius:20px;font-size:11px;font-weight:500;margin:3px}
.t-match{background:rgba(63,185,80,.12);color:#3fb950;border:1px solid rgba(63,185,80,.25)}
.t-miss {background:rgba(210,153,34,.12);color:#d29922;border:1px solid rgba(210,153,34,.25)}
.sug{padding:10px 14px;border-radius:8px;font-size:13px;line-height:1.5;margin-bottom:8px;border-left:3px solid}
.sug.h{background:rgba(248,81,73,.07);border-color:var(--red);color:#ffb3af}
.sug.m{background:rgba(210,153,34,.07);border-color:var(--yellow);color:#f0c040}
.sug.l{background:rgba(63,185,80,.07);border-color:var(--green);color:#7ee787}

/* download */
.dl-row{display:flex;gap:12px;justify-content:center;flex-wrap:wrap;margin-top:8px}
.dl-btn{padding:12px 26px;border-radius:8px;font-size:13.5px;font-weight:600;
  border:none;cursor:pointer;font-family:'Inter',sans-serif;
  transition:opacity .2s,transform .1s;text-decoration:none;display:inline-block;color:#fff}
.dl-btn:hover{opacity:.85;transform:translateY(-1px)}
.dl-docx{background:var(--blue)}
.dl-pdf {background:#e63946}

@media(max-width:768px){.grid{grid-template-columns:1fr}.stats{grid-template-columns:1fr 1fr}}
</style>
</head>
<body>
<div class="wrap">
  <div class="hdr">
    <div class="badge">AI Powered · ATS Optimized</div>
    <h1>Resume Optimizer</h1>
    <p>Match your resume to any job — improve keywords, metrics &amp; ATS score</p>
  </div>

  <div class="grid">
    <div class="card">
      <div class="card-lbl">📄 Your Resume</div>
      <textarea id="resumeText" placeholder="Paste your resume text here...&#10;&#10;Jane Doe&#10;jane@email.com | linkedin.com/in/jane&#10;&#10;WORK EXPERIENCE&#10;Software Engineer | Company | 2021–Present&#10;• Reduced API latency by 40%&#10;• Led migration saving $80,000 annually"></textarea>
    </div>
    <div class="card">
      <div class="card-lbl">💼 Job Description</div>
      <textarea id="jobText" placeholder="Paste the job description here...&#10;&#10;Senior Software Engineer&#10;We are looking for an engineer experienced in Python, AWS, Docker, Kubernetes, CI/CD..."></textarea>
    </div>
  </div>

  <button class="btn" id="analyzeBtn" onclick="analyze()">Analyze &amp; Optimize →</button>
  <div class="spin-wrap" id="spinner"><span class="spin-icon"></span> Analyzing your resume…</div>

  <div id="results" style="margin-top:20px"></div>
</div>

<script>
async function analyze(){
  const resume=document.getElementById('resumeText').value.trim();
  const job=document.getElementById('jobText').value.trim();
  if(!resume||!job){alert('Please paste both your resume and the job description.');return;}
  const btn=document.getElementById('analyzeBtn');
  const sp=document.getElementById('spinner');
  btn.disabled=true;btn.textContent='Analyzing…';
  sp.classList.add('on');
  document.getElementById('results').style.display='none';
  const fd=new FormData();
  fd.append('resume_text',resume);fd.append('job_text',job);
  try{
    const res=await fetch('/analyze',{method:'POST',body:fd});
    const d=await res.json();
    if(!res.ok){alert(d.error||'Error');return;}
    render(d);
  }catch(e){alert('Error: '+e.message);}
  finally{btn.disabled=false;btn.textContent='Analyze & Optimize →';sp.classList.remove('on');}
}

function render(d){
  const vlabel={excellent:'🏆 Excellent Match',good:'✅ Good Match',
    moderate:'⚠️ Moderate Match',poor:'❌ Needs Improvement'}[d.verdict];

  const sugHtml=d.impact_suggestions.map((s,i)=>`
    <div class="sug ${i<2?'h':i<3?'m':'l'}">${i+1}. ${s}</div>`).join('');
  const kwHtml=d.keyword_suggestions.map((s,i)=>`
    <div class="sug ${i===0?'h':'m'}">${s}</div>`).join('');
  const mTags=d.matching_keywords.map(k=>`<span class="tag t-match">${k}</span>`).join('');
  const xTags=d.missing_keywords.map(k=>`<span class="tag t-miss">${k}</span>`).join('');

  document.getElementById('results').innerHTML=`
    <div class="hero">
      <div class="score-num c-${d.verdict}">${d.match_score}%</div>
      <div class="score-lbl c-${d.verdict}">${vlabel}</div>
      <div class="pill">${d.verdict_message}</div>
    </div>
    <div class="stats">
      <div class="stat"><div class="stat-v">${d.total_matching}</div><div class="stat-l">Keywords Matched</div></div>
      <div class="stat"><div class="stat-v">${d.total_missing}</div><div class="stat-l">Skills to Add</div></div>
      <div class="stat"><div class="stat-v">${d.impact_score}</div><div class="stat-l">Impact Score /100</div></div>
      <div class="stat"><div class="stat-v">${d.total_impacts}</div><div class="stat-l">Achievements Found</div></div>
    </div>
    ${sugHtml||kwHtml?`<div class="panel"><div class="panel-t">Priority Improvements</div>${sugHtml}${kwHtml}</div>`:''}
    <div class="panel">
      <div class="panel-t">✅ Keywords Matched (${d.matching_keywords.length})</div>
      ${mTags||'<span style="color:var(--muted)">None detected</span>'}
    </div>
    <div class="panel">
      <div class="panel-t">✨ Keywords Added (${d.missing_keywords.length})</div>
      ${xTags||'<span style="color:var(--muted)">None — great coverage!</span>'}
    </div>
    <div class="panel" style="text-align:center">
      <div class="panel-t" style="text-align:center">Download Optimized Resume</div>
      <p style="color:var(--muted);font-size:12px;margin-bottom:14px">
        Candidate: <strong style="color:var(--text)">${d.name}</strong>
      </p>
      <div class="dl-row">
        <a class="dl-btn dl-docx" href="/download/${encodeURIComponent(d.filename)}">📄 Download Word (.docx)</a>
        <a class="dl-btn dl-pdf"  href="/download/${encodeURIComponent(d.pdf_filename)}">📕 Download PDF</a>
      </div>
    </div>`;
  document.getElementById('results').style.display='block';
  document.getElementById('results').scrollIntoView({behavior:'smooth',block:'start'});
}
</script>
</body>
</html>'''


@app.route('/')
def index():
    return render_template_string(HTML)

@app.route('/analyze', methods=['POST'])
def analyze():
    try:
        r = request.form.get('resume_text','')
        j = request.form.get('job_text','')
        if not r or not j:
            return jsonify({'error':'Provide both resume and job description'}), 400
        return jsonify(matcher.analyze(r, j))
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download/<path:filename>')
def download_file(filename):
    path = os.path.join(app.config['OPTIMIZED_FOLDER'], filename)
    if not os.path.exists(path):
        return jsonify({'error':'File not found'}), 404
    mime = ('application/pdf' if filename.endswith('.pdf') else
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return send_file(path, as_attachment=True, download_name=filename, mimetype=mime)

if __name__ == '__main__':
    print("\n  Resume Optimizer  →  http://localhost:5000\n")
    app.run(debug=True, host='127.0.0.1', port=5000)